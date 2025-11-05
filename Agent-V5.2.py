import sys
import os
import re
import threading
import time
import docx
import openpyxl
import xml.etree.ElementTree as ET
import pytesseract
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image, ImageFilter, ImageOps, ImageEnhance
from PyQt6.QtCore import pyqtSignal, Qt
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QTextEdit, QLineEdit,
    QPushButton, QLabel, QFileDialog, QMessageBox, QComboBox
)
from PyQt6.QtGui import QFont, QTextCursor
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from collections import OrderedDict
import pandas as pd
import sys
import os

# === Paths for bundled Tesseract and Poppler ===
if getattr(sys, 'frozen', False):  # Running as EXE
    base_path = sys._MEIPASS  # PyInstaller temp dir
else:
    base_path = os.path.dirname(os.path.abspath(__file__))

# Paths inside the build
TESSERACT_PATH = os.path.join(base_path, "Tesseract-OCR", "tesseract.exe")
POPPLER_PATH = os.path.join(base_path, "poppler", "Library", "bin")

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
# Optional: set poppler_path if convert_from_path requires it (Windows)
#POPPLER_PATH = r"C:\poppler\Library\bin"  # adjust if needed

# === Tesseract path ===
#pytesseract.pytesseract.tesseract_cmd = r"C:\Users\omoham108\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

# === Ollama/DeepSeek config ===
OLLAMA_BASE_URL = "http://10.11.3.181:11434"
# Models available for selection
AVAILABLE_MODELS = {
    "DeepSeek R1 70B": "deepseek-r1:70b",
    "DeepSeek Coder 33B": "deepseek-coder:33b",
    "Qwen3": "Qwen3"

}
DEFAULT_MODEL = "deepseek-r1:70b"

# Supported file extensions for multiple file upload
SUPPORTED_EXTENSIONS = [
    "*.png", "*.jpg", "*.jpeg", "*.svg",  # Images
    "*.pdf",  # PDF
    "*.docx",  # Word
    "*.vsdx",  # Visio
    "*.py", "*.c", "*.cpp", "*.h", "*.hpp", "*.java", "*.js", "*.ts",  # Code files
    "*.html", "*.css", "*.xml", "*.json", "*.yaml", "*.yml",  # Web/Config
    "*.txt", "*.md", "*.rst",  # Text files
    "*.csv", "*.xlsx", "*.xls"  # Data files
]

# ----------------------
# Memory Manager
# ----------------------
class MemoryManager:
    def __init__(self):
        self.memory = []
        self.vectorizer = TfidfVectorizer()

    def add(self, user_input, model_output):
        self.memory.append(f"User: {user_input}\nAssistant: {model_output}")

    def recall(self, query, top_k=3):
        if not self.memory:
            return ""
        docs = self.memory + [query]
        try:
            vectors = self.vectorizer.fit_transform(docs)
            similarities = cosine_similarity(vectors[-1:], vectors[:-1]).flatten()
            top_indices = similarities.argsort()[-top_k:][::-1]
            return "\n\n".join([self.memory[i] for i in top_indices])
        except Exception:
            return ""

# ----------------------
# LLM caller (lazy load)
# ----------------------
_llm_client = None
_llm_client_lock = threading.Lock()
_current_model = DEFAULT_MODEL

def get_llm(model_name=None):
    global _llm_client, _current_model
    with _llm_client_lock:
        if model_name and model_name != _current_model:
            _llm_client = None  # Force recreation if model changed
            _current_model = model_name
            
        if _llm_client is None:
            try:
                from langchain_ollama import OllamaLLM
                _llm_client = OllamaLLM(model=_current_model, base_url=OLLAMA_BASE_URL)
            except Exception as e:
                _llm_client = e
        return _llm_client

def call_langchain_model(prompt: str, model_name=None):
    client = get_llm(model_name)
    if isinstance(client, Exception):
        return f"❌ Error initializing Ollama client: {client}"
    try:
        response = client.invoke(prompt)
        return response.strip() if response else "⚠️ Empty response from model."
    except Exception as e:
        return f"❌ LangChain/Ollama error: {e}"

# ----------------------
# File save functions
# ----------------------
def save_response_to_txt(text, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path

def save_response_to_docx(text, path):
    doc = docx.Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)
    return path

def save_response_to_pdf(text, path):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 40
    text_object = c.beginText(40, y)
    text_object.setFont("Helvetica", 10)
    for line in text.splitlines():
        text_object.textLine(line)
        y -= 12
        if y < 40:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(40, height - 40)
            text_object.setFont("Helvetica", 10)
            y = height - 40
    c.drawText(text_object)
    c.save()
    return path

# ----------------------
# Image preprocessing (always ON)
# ----------------------
def preprocess_image_for_ocr(pil_img: Image.Image):
    img = pil_img.convert("RGB")
    img = ImageOps.exif_transpose(img)
    img = ImageOps.grayscale(img)
    img = ImageEnhance.Contrast(img).enhance(1.6)
    img = ImageEnhance.Sharpness(img).enhance(1.2)
    img = img.filter(ImageFilter.MedianFilter(size=3))
    return img

# ----------------------
# File processing functions
# ----------------------
def extract_text_from_file(file_path):
    """Extract text from various file types"""
    file_ext = os.path.splitext(file_path)[1].lower()
    text_content = ""
    
    try:
        if file_ext == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text_content += (page.extract_text() or "") + "\n"
            if not text_content.strip():
                imgs = convert_from_path(file_path, poppler_path=POPPLER_PATH if os.path.exists(POPPLER_PATH) else None)
                for img in imgs:
                    img = preprocess_image_for_ocr(img)
                    text_content += pytesseract.image_to_string(img)
                    
        elif file_ext in ['.docx']:
            doc = docx.Document(file_path)
            text_content = "\n".join([p.text for p in doc.paragraphs])
            
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            img = Image.open(file_path)
            img = preprocess_image_for_ocr(img)
            text_content = pytesseract.image_to_string(img)
            
        elif file_ext == '.svg':
            tree = ET.parse(file_path)
            root = tree.getroot()
            texts = [elem.text.strip() for elem in root.iter() if elem.text and elem.text.strip()]
            text_content = "\n".join(texts)
            
        elif file_ext == '.vsdx':
            from vsdx import VisioFile
            vf = VisioFile(file_path)
            texts = []
            for page in vf.pages:
                for shape in page.shapes:
                    if getattr(shape, "text", None):
                        texts.append(shape.text.strip())
            text_content = "\n".join(texts)

        elif file_ext in ['.xlsx', '.xls']:
            # Try pandas first
            try:
                xls = pd.ExcelFile(file_path, engine='openpyxl')
                text_rows = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str, engine='openpyxl')
                    text_rows.append(f"--- Sheet: {sheet_name} ---")
                    for row in df.values:
                        text_rows.append(" | ".join([str(cell) for cell in row if pd.notna(cell)]))
                text_content = "\n".join(text_rows)
            except Exception as e1:
                # Fallback: use xlwings
                try:
                    import xlwings as xw
                    wb = xw.Book(file_path)
                    text_rows = []
                    for sheet in wb.sheets:
                        text_rows.append(f"--- Sheet: {sheet.name} ---")
                        data = sheet.used_range.value
                        if data:
                            for row in data:
                                text_rows.append(" | ".join([str(cell) for cell in row if cell is not None]))
                    wb.close()
                    text_content = "\n".join(text_rows)
                except Exception as e2:
                    text_content = f"❌ Error reading Excel: pandas error: {e1}; xlwings fallback error: {e2}"



            
        else:  # Text-based files (code, txt, etc.)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
                
    except Exception as e:
        return f"❌ Error reading {file_path}: {str(e)}"
    
    return text_content.strip()

# ----------------------
# Main GUI
# ----------------------
class AgentApp(QWidget):
    # Signals to ensure UI updates occur on the main thread
    response_ready = pyqtSignal(str)               # for model responses
    system_message = pyqtSignal(str)               # for system/notice messages
    status_update = pyqtSignal(str, str)           # status text and stylesheet (both strings)

    def __init__(self):
        super().__init__()
        self.memory = MemoryManager()
        self.last_output = ""
        # Use OrderedDict keyed by full path to avoid basename collisions and preserve order
        self.uploaded_files = OrderedDict()
        self.selected_model = DEFAULT_MODEL

        # connect signals to slots (main thread)
        self.response_ready.connect(self._on_response_ready)
        self.system_message.connect(self._on_system_message)
        self.status_update.connect(self._on_status_update)

        self.setWindowTitle("Ragura AI Agent")
        self.setGeometry(100, 100, 1200, 800)
        self.setup_ui()

        # Start connection check in background
        threading.Thread(target=self.check_ollama_connection, daemon=True).start()

    def setup_ui(self):
        # Main layout
        main_layout = QHBoxLayout()
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Left panel (Chat area - 75%)
        chat_container = QWidget()
        chat_container.setObjectName("chatContainer")
        chat_layout = QVBoxLayout(chat_container)
        chat_layout.setContentsMargins(20, 20, 20, 20)
        chat_layout.setSpacing(15)

        # Chat header
        header = QLabel("Ragura AI Assistant")
        header.setObjectName("chatHeader")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header.setMaximumHeight(60)
        chat_layout.addWidget(header)

        # Chat display area
        self.chat_display = QTextEdit()
        self.chat_display.setObjectName("chatDisplay")
        self.chat_display.setReadOnly(True)
        self.chat_display.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        chat_layout.addWidget(self.chat_display)

        # Input area
        input_container = QWidget()
        input_container.setMaximumHeight(80)
        input_layout = QHBoxLayout(input_container)
        input_layout.setContentsMargins(0, 10, 0, 0)

        self.input_box = QLineEdit()
        self.input_box.setObjectName("inputBox")
        self.input_box.setPlaceholderText("Type your message here...")
        self.input_box.returnPressed.connect(self.send_message)

        self.send_button = QPushButton("Send")
        self.send_button.setObjectName("sendButton")
        self.send_button.clicked.connect(self.send_message)
        self.send_button.setFixedSize(80, 40)

        input_layout.addWidget(self.input_box)
        input_layout.addWidget(self.send_button)
        chat_layout.addWidget(input_container)

        # Right panel (Controls - 25%)
        controls_container = QWidget()
        controls_container.setObjectName("controlsContainer")
        controls_container.setFixedWidth(300)
        controls_layout = QVBoxLayout(controls_container)
        controls_layout.setContentsMargins(15, 20, 15, 20)
        controls_layout.setSpacing(15)

        # Model Selection
        controls_layout.addWidget(QLabel("Select Model"))
        self.model_combo = QComboBox()
        self.model_combo.setObjectName("modelCombo")
        for display_name in AVAILABLE_MODELS.keys():
            self.model_combo.addItem(display_name)
        self.model_combo.currentTextChanged.connect(self.on_model_changed)
        controls_layout.addWidget(self.model_combo)

        # Status
        self.status_label = QLabel("Ready")
        self.status_label.setFixedSize(240, 50)  # fixed width and height
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setObjectName("statusLabel")
        controls_layout.addWidget(QLabel("Agent Status"))
        controls_layout.addWidget(self.status_label)

        # File upload section
        controls_layout.addWidget(QLabel("Upload Files"))
        
        self.upload_multiple = QPushButton("📁 Upload Multiple Files")
        self.upload_folder = QPushButton("📂 Upload Folder")
        #self.upload_flowchart = QPushButton("📊 Upload Flowchart Image")
        #self.upload_pdf = QPushButton("📄 Upload PDF")
        #self.upload_word = QPushButton("📝 Upload Word")
        #self.upload_visio = QPushButton("🔷 Upload Visio")
        self.clear_uploads_btn = QPushButton("🧹 Clear Uploads")  # Added: clear uploads

        #for btn in [self.upload_multiple, self.upload_folder, self.upload_flowchart, 
                 #  self.upload_pdf, self.upload_word, self.upload_visio, self.clear_uploads_btn]:
        for btn in [self.upload_multiple, self.upload_folder, self.clear_uploads_btn]:
            btn.setObjectName("uploadButton")
            btn.setFixedHeight(40)

        self.upload_multiple.clicked.connect(self.upload_multiple_files)
        self.upload_folder.clicked.connect(self.upload_folder_files)
        #self.upload_flowchart.clicked.connect(lambda: self.upload_file("Flowchart", ["*.png", "*.jpg", "*.jpeg", "*.svg"]))
        #self.upload_pdf.clicked.connect(lambda: self.upload_file("PDF", ["*.pdf"]))
        #self.upload_word.clicked.connect(lambda: self.upload_file("Word", ["*.docx"]))
        #self.upload_visio.clicked.connect(lambda: self.upload_file("Visio", ["*.vsdx"]))
        self.clear_uploads_btn.clicked.connect(self.clear_uploaded_files)  # hook clear

        controls_layout.addWidget(self.upload_multiple)
        controls_layout.addWidget(self.upload_folder)
        #controls_layout.addWidget(self.upload_flowchart)
        #controls_layout.addWidget(self.upload_pdf)
        #controls_layout.addWidget(self.upload_word)
        #controls_layout.addWidget(self.upload_visio)
        controls_layout.addWidget(self.clear_uploads_btn)  # show clear button

        # Export section
        controls_layout.addWidget(QLabel("Export Response"))
        self.export_pdf_btn = QPushButton("📥 Save as PDF")
        self.export_docx_btn = QPushButton("📥 Save as Word")
        self.export_txt_btn = QPushButton("📥 Save as TXT")

        for btn in [self.export_pdf_btn, self.export_docx_btn, self.export_txt_btn]:
            btn.setObjectName("exportButton")
            btn.setFixedHeight(40)

        self.export_pdf_btn.clicked.connect(lambda: self.export_last_response("pdf"))
        self.export_docx_btn.clicked.connect(lambda: self.export_last_response("docx"))
        self.export_txt_btn.clicked.connect(lambda: self.export_last_response("txt"))

        controls_layout.addWidget(self.export_pdf_btn)
        controls_layout.addWidget(self.export_docx_btn)
        controls_layout.addWidget(self.export_txt_btn)

        # Connection test
        self.conn_btn = QPushButton("🔌 Test Connection")
        self.conn_btn.setObjectName("connectionButton")
        self.conn_btn.setFixedHeight(40)
        self.conn_btn.clicked.connect(lambda: threading.Thread(target=self.check_ollama_connection, daemon=True).start())
        controls_layout.addWidget(self.conn_btn)

        controls_layout.addStretch()

        # Add panels to main layout
        main_layout.addWidget(chat_container)
        main_layout.addWidget(controls_container)

        self.setLayout(main_layout)
        self.apply_styles()

        # Initial system message via signal (main thread)
        self.system_message.emit("🤖 Welcome to Ragura AI Assistant! I'm ready to help you with document analysis, questions, and more.")

    def apply_styles(self):
        self.setStyleSheet("""
            QWidget {
                font-family: 'Segoe UI', Arial, sans-serif;
                background: #0f1720;
            }
            #chatContainer {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #0f1720, stop:1 #1e293b);
                border-right: 1px solid #334155;
            }
            /* ✅ Remove black background under input area */
            QWidget > QWidget > QWidget {
                background: transparent;
                border: none;
            }
            #modelCombo {
                background: #1e293b;
                color: #e2e8f0;
                border: 2px solid #475569;
                border-radius: 8px;
                padding: 8px 12px;
                font-size: 14px;
                min-height: 20px;
            }
            #modelCombo:focus {
                border-color: #6366f1;
            }
            #modelCombo QAbstractItemView {
                background: #1e293b;
                color: #e2e8f0;
                border: 1px solid #475569;
                border-radius: 8px;
                selection-background-color: #6366f1;
            }

            #chatHeader {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1e40af, stop:1 #7c3aed);
                color: white;
                font-size: 20px;
                font-weight: bold;
                padding: 12px;
                border-radius: 10px;
                border: 1px solid #4f46e5;
            }
            #chatDisplay {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #0f1720, stop:1 #1e293b);
                color: #e2e8f0;
                border: 2px solid #334155;
                border-radius: 15px;
                padding: 15px;
                font-size: 14px;
                line-height: 1.4;
                selection-background-color: transparent;
            }
            #chatDisplay QScrollBar:vertical {
                background: #1e293b;
                width: 12px;
                border-radius: 6px;
            }
            #chatDisplay QScrollBar::handle:vertical {
                background: #4f46e5;
                border-radius: 6px;
                min-height: 20px;
            }
            #chatDisplay QScrollBar::handle:vertical:hover {
                background: #6366f1;
            }
            #inputBox {
                background: #1e293b;
                color: #e2e8f0;
                border: 2px solid #475569;
                border-radius: 20px;
                padding: 12px 20px;
                font-size: 14px;
                margin-right: 10px;
            }
            #inputBox:focus {
                border-color: #6366f1;
                background: #1e293b;
            }
            #sendButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #6366f1, stop:1 #4f46e5);
                color: white;
                border: none;
                border-radius: 20px;
                font-weight: bold;
                font-size: 14px;
            }
            #sendButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #4f46e5, stop:1 #4338ca);
            }
            #sendButton:pressed {
                background: #3730a3;
            }
            #controlsContainer {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #1e293b, stop:1 #0f172a);
            }
            #statusLabel {
                color: #10b981;
                font-size: 16px;
                font-weight: bold;
                padding: 8px;
                background: rgba(16, 185, 129, 0.1);
                border-radius: 8px;
                border: 1px solid #10b981;
            }
            #uploadButton, #exportButton, #connectionButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #475569, stop:1 #334155);
                color: #e2e8f0;
                border: 1px solid #64748b;
                border-radius: 8px;
                font-size: 13px;
                text-align: left;
                padding-left: 15px;
            }
            #uploadButton:hover, #exportButton:hover, #connectionButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #6366f1, stop:1 #4f46e5);
                border-color: #818cf8;
            }
            #uploadButton:pressed, #exportButton:pressed, #connectionButton:pressed {
                background: #3730a3;
            }
            QLabel {
                color: #94a3b8;
                font-weight: bold;
                margin-bottom: 5px;
            }
        """)

    # ---- Model Selection ----
    def on_model_changed(self, display_name):
        self.selected_model = AVAILABLE_MODELS.get(display_name, DEFAULT_MODEL)
        self.system_message.emit(f"🔄 Model changed to: {display_name}")
        # Reset LLM client to force recreation with new model
        global _llm_client
        _llm_client = None

    # ---- Slots for signals (main thread) ----
    def _on_response_ready(self, text: str):
        self.append_model_message(text)
        # update status to ready
        style = "color: #10b981; background: rgba(16, 185, 129, 0.1); border: 1px solid #10b981;"
        self.status_update.emit("Ready", style)

    def _on_system_message(self, text: str):
        self.append_system_message(text)

    def _on_status_update(self, text: str, style: str):
        self.status_label.setText(text)
        try:
            self.status_label.setStyleSheet(style)
        except Exception:
            # fallback: set plain text color if style invalid
            pass

    # ---- Helper append methods (main thread) ----
    def append_system_message(self, text):
        html = f'''
        <table width="100%" cellpadding="0" cellspacing="0" style="margin: 15px 0;">
            <tr>
                <td align="center">
                    <div style="
                        background: rgba(99, 102, 241, 0.15);
                        color: #c7d2fe;
                        padding: 10px 20px;
                        border-radius: 18px;
                        border: 1px solid rgba(99, 102, 241, 0.3);
                        font-size: 13px;
                        font-style: italic;
                        max-width: 80%;
                        display: inline-block;">
                        {self._escape_html(text)}
                    </div>
                </td>
            </tr>
        </table>
        '''
        self._append_html(html)

    def append_user_message(self, text):
        html = f'''
        <table width="100%" cellpadding="0" cellspacing="0" style="margin: 15px 0;">
            <tr>
                <td align="right">
                    <table cellpadding="8" cellspacing="0" style="
                        background-color: #6366f1;
                        color: white;
                        border-radius: 12px 12px 5px 12px;
                        max-width: 70%;
                        border: 1px solid rgba(255, 255, 255, 0.1);
                        word-wrap: break-word;
                    ">
                        <tr>
                            <td>{self._escape_html(text)}</td>
                        </tr>
                    </table>
                    <div style="color: #94a3b8; font-size: 11px; margin-top: 5px; margin-right: 5px;">
                        You
                    </div>
                </td>
            </tr>
        </table>
        '''
        self._append_html(html)


    def append_model_message(self, text):
        html = f'''
        <table width="100%" cellpadding="0" cellspacing="0" style="margin: 15px 0;">
            <tr>
                <td align="left">
                    <table cellpadding="8" cellspacing="0" style="
                        background-color: #374151;
                        color: #f1f5f9;
                        border-radius: 12px 12px 12px 5px;
                        max-width: 70%;
                        border: 1px solid rgba(255, 255, 255, 0.05);
                        word-wrap: break-word;
                    ">
                        <tr>
                            <td>{self._escape_html(text)}</td>
                        </tr>
                    </table>
                    <div style="color: #94a3b8; font-size: 11px; margin-top: 5px; margin-left: 5px;">
                         Agent
                    </div>
                </td>
            </tr>
        </table>
        '''
        self._append_html(html)


    def _append_html(self, html):
        # All UI changes come through signals -> main thread -> here
        self.chat_display.append(html)
        self.chat_display.moveCursor(QTextCursor.MoveOperation.End)
        QApplication.processEvents()

    def _escape_html(self, s):
        return (s.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("\n", "<br>"))

    # ---- File upload methods ----
    def upload_file(self, file_type, extensions):
        file_filter = f"{file_type} Files ({' '.join(extensions)})"
        file_path, _ = QFileDialog.getOpenFileName(self, f"Select {file_type}", "", file_filter)
        if not file_path:
            return
        self.process_single_file(file_path)

    def upload_multiple_files(self):
        file_filter = f"All Supported Files ({' '.join(SUPPORTED_EXTENSIONS)})"
        file_paths, _ = QFileDialog.getOpenFileNames(self, "Select Files", "", file_filter)
        if not file_paths:
            return
        
        self.system_message.emit(f"📁 Processing {len(file_paths)} files...")
        for file_path in file_paths:
            self.process_single_file(file_path)

    def upload_folder_files(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder")
        if not folder_path:
            return
        
        supported_files = []
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                if any(file_path.lower().endswith(ext.replace('*', '')) for ext in SUPPORTED_EXTENSIONS):
                    supported_files.append(file_path)
        
        if not supported_files:
            self.system_message.emit("❌ No supported files found in the selected folder.")
            return
            
        self.system_message.emit(f"📂 Processing {len(supported_files)} files from folder...")
        for file_path in supported_files:
            self.process_single_file(file_path)

    def process_single_file(self, file_path): 
        file_name = os.path.basename(file_path)
        try:
            text_content = extract_text_from_file(file_path)
            if text_content.startswith("❌ Error"):
                self.system_message.emit(f"❌ {text_content}")
            else:
                # Key by full path to avoid collisions; preserve order
                prev_keys_with_same_basename = [k for k in self.uploaded_files.keys() if os.path.basename(k) == file_name]
                self.uploaded_files[file_path] = text_content
                file_size = len(text_content)
                # Inform user; warn if same basename exists from another path
                if prev_keys_with_same_basename:
                    self.system_message.emit(f"📎 Uploaded: {file_name} ({file_size} chars) — note: another file with the same name exists (different path).")
                else:
                    self.system_message.emit(f"📎 Uploaded: {file_name} ({file_size} chars)")
        except Exception as e:
            self.system_message.emit(f"❌ Error processing {file_name}: {str(e)}")

    def clear_uploaded_files(self):
        count = len(self.uploaded_files)
        self.uploaded_files.clear()
        self.system_message.emit(f"🧹 Cleared {count} uploaded file(s).")

    # ---- Send message ----
    def send_message(self):
        user_text = self.input_box.text().strip()
        if not user_text:
            return
        self.append_user_message(user_text)
        self.input_box.clear()
        self.status_update.emit("Thinking...", "color: #f59e0b; background: rgba(245, 158, 11, 0.1); border: 1px solid #f59e0b;")

        context = self.memory.recall(user_text)
        if self.uploaded_files:
            # Build files text in upload order; show only filename, but keep content mapped by full path
            files_text = []
            for fullpath, content in self.uploaded_files.items():
                name = os.path.basename(fullpath)
                # limit per file to first N chars to avoid overly long prompts
                files_text.append(f"--- {name} ---\n{content}")
            context += f"\n\n[Uploaded Files Content]\n" + "\n\n".join(files_text)

        full_prompt = f"{context}\n\nUser: {user_text}\nAssistant:"
        threading.Thread(target=self.process_message, args=(user_text, full_prompt), daemon=True).start()

    def process_message(self, user_text, prompt):
        response = call_langchain_model(prompt, self.selected_model)
        self.memory.add(user_text, response)
        self.last_output = response
        # Emit model response via signal so UI updates on main thread
        self.response_ready.emit(response)

    # ---- Export ----
    def export_last_response(self, fmt):
        if not self.last_output:
            QMessageBox.information(self, "No response", "No response to save yet.")
            return
        name = f"DeepSeek_Response_{int(time.time())}"
        try:
            if fmt == "pdf":
                path, _ = QFileDialog.getSaveFileName(self, "Save PDF", f"{name}.pdf", "PDF Files (*.pdf)")
                if path:
                    try:
                        save_response_to_pdf(self.last_output, path)
                        self.system_message.emit(f"✅ Response saved as PDF: {os.path.basename(path)}")
                    except Exception as e:
                        QMessageBox.critical(self, "Save PDF Error", f"Failed to save PDF:\n{e}")
            elif fmt == "docx":
                path, _ = QFileDialog.getSaveFileName(self, "Save Word", f"{name}.docx", "Word Files (*.docx)")
                if path:
                    try:
                        save_response_to_docx(self.last_output, path)
                        self.system_message.emit(f"✅ Response saved as Word: {os.path.basename(path)}")
                    except Exception as e:
                        QMessageBox.critical(self, "Save DOCX Error", f"Failed to save Word file:\n{e}")
            elif fmt == "txt":
                path, _ = QFileDialog.getSaveFileName(self, "Save TXT", f"{name}.txt", "Text Files (*.txt)")
                if path:
                    try:
                        save_response_to_txt(self.last_output, path)
                        self.system_message.emit(f"✅ Response saved as TXT: {os.path.basename(path)}")
                    except Exception as e:
                        QMessageBox.critical(self, "Save TXT Error", f"Failed to save TXT file:\n{e}")
        except Exception as e:
            # Catch any unexpected error from QFileDialog itself
            QMessageBox.critical(self, "Export Error", f"An unexpected error occurred:\n{e}")

    # ---- Connection Test ----
    def check_ollama_connection(self):
        # Runs in background thread; use signals for UI update
        try:
            import requests
            url = OLLAMA_BASE_URL.rstrip("/") + "/api/ping"
            r = requests.get(url, timeout=3)
            if r.ok:
                self.system_message.emit(f"✅ Connected to Ollama server at {OLLAMA_BASE_URL}")
                self.status_update.emit("Connected", "color: #10b981; background: rgba(16, 185, 129, 0.1); border: 1px solid #10b981;")
            else:
                self.system_message.emit(f"⚠️ Server reachable but status {r.status_code}")
                self.status_update.emit("Server Error", "color: #ef4444; background: rgba(239, 68, 68, 0.1); border: 1px solid #ef4444;")
        except Exception as e:
            self.system_message.emit(f"❌ Cannot reach Ollama server: {e}")
            self.status_update.emit("Disconnected", "color: #ef4444; background: rgba(239, 68, 68, 0.1); border: 1px solid #ef4444;")
        finally:
            # re-enable connection button on main thread
            try:
                self.conn_btn.setEnabled(True)
            except Exception:
                pass

# ---- Run App ----
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = AgentApp()
    window.show()
    sys.exit(app.exec())
